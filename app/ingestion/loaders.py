from __future__ import annotations

import logging
from pathlib import Path
from typing import List

import pandas as pd
from pydantic import BaseModel

from app.core.exceptions import IngestionError

logger = logging.getLogger(__name__)


class LoadedDocument(BaseModel):
    """Normalised output of every loader — one object per file."""

    text_content: str
    tables: List[object] = []   # List[pd.DataFrame], stored as generic objects for Pydantic
    metadata: dict

    class Config:
        arbitrary_types_allowed = True


def _load_pdf(path: Path) -> LoadedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    logger.info(f"PDF loaded: {path.name} — {len(reader.pages)} pages, {len(full_text)} chars")
    return LoadedDocument(
        text_content=full_text,
        metadata={"source": str(path), "type": "pdf", "pages": len(reader.pages)},
    )


def _load_docx(path: Path) -> LoadedDocument:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    logger.info(f"DOCX loaded: {path.name} — {len(paragraphs)} paragraphs")
    return LoadedDocument(
        text_content=full_text,
        metadata={"source": str(path), "type": "docx"},
    )


def _load_excel(path: Path) -> LoadedDocument:
    xl = pd.ExcelFile(str(path))
    tables: list[pd.DataFrame] = []
    text_parts: list[str] = []

    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        df = df.dropna(how="all").dropna(axis=1, how="all")
        if df.empty:
            continue
        tables.append(df)
        # Also convert to text so the vector store can index it
        text_parts.append(f"[Sheet: {sheet}]\n{df.to_string(index=False)}")

    full_text = "\n\n".join(text_parts)
    logger.info(f"Excel loaded: {path.name} — {len(tables)} sheets, {sum(len(t) for t in tables)} rows total")
    return LoadedDocument(
        text_content=full_text,
        tables=tables,
        metadata={"source": str(path), "type": "excel", "sheets": xl.sheet_names},
    )


def _load_text(path: Path) -> LoadedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    logger.info(f"Text loaded: {path.name} — {len(text)} chars")
    return LoadedDocument(
        text_content=text,
        metadata={"source": str(path), "type": "text"},
    )


def _load_html(path: Path) -> LoadedDocument:
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        # Skip tags that never contain human-readable content
        SKIP_TAGS = {"script", "style", "head", "meta", "link", "ix:hidden", "ix:header"}

        def __init__(self):
            super().__init__()
            self._skip_depth = 0
            self.chunks: list[str] = []

        def handle_starttag(self, tag, attrs):
            if tag.lower() in self.SKIP_TAGS:
                self._skip_depth += 1

        def handle_endtag(self, tag):
            if tag.lower() in self.SKIP_TAGS and self._skip_depth > 0:
                self._skip_depth -= 1

        def handle_data(self, data):
            if self._skip_depth == 0:
                stripped = data.strip()
                if stripped:
                    self.chunks.append(stripped)

    html = path.read_text(encoding="utf-8", errors="replace")
    extractor = _TextExtractor()
    extractor.feed(html)
    text = "\n".join(extractor.chunks)
    logger.info(f"HTML loaded: {path.name} — {len(text)} chars")
    return LoadedDocument(
        text_content=text,
        metadata={"source": str(path), "type": "html"},
    )


class DocumentLoader:
    """Dispatch file loading based on extension."""

    _DISPATCH = {
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".xlsx": _load_excel,
        ".xls": _load_excel,
        ".txt": _load_text,
        ".md": _load_text,
        ".htm": _load_html,
        ".html": _load_html,
    }

    def load(self, path: str | Path, display_name: str | None = None) -> LoadedDocument:
        path = Path(path)
        if not path.exists():
            raise IngestionError(f"File not found: {path}")

        loader_fn = self._DISPATCH.get(path.suffix.lower())
        if loader_fn is None:
            raise IngestionError(f"Unsupported file type: {path.suffix}")

        try:
            doc = loader_fn(path)
        except IngestionError:
            raise
        except Exception as e:
            raise IngestionError(f"Failed to load {path.name}: {e}") from e

        # Override the temp-file source with the human-readable filename so
        # chunk metadata propagates the real filename through ChromaDB → citations → map.
        if display_name:
            doc.metadata["source"] = display_name
        return doc
